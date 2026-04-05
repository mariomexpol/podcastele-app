import streamlit as st
import requests
import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(page_title="PodcastELE Pro - Fix 2024", layout="wide", page_icon="🎙️")

# --- FUNCIONES DE SOPORTE ---
def limpiar_texto(texto):
    return texto.replace('\\_', '_').replace('\\', '')

def generar_docx_podcast(texto_ia, escuela, profe, tema, nivel, logo_file=None):
    doc = Document()
    section = doc.sections[0]
    header = section.header
    htxt = header.paragraphs[0]
    
    if logo_file:
        try:
            run_logo = htxt.add_run()
            run_logo.add_picture(logo_file, width=Inches(1.2))
            htxt.alignment = WD_ALIGN_PARAGRAPH.LEFT
        except: pass
    
    info_h = header.add_paragraph(f"{escuela}\nMaterial ELE - Nivel {nivel}\nProf. {profe}")
    info_h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph() 
    titulo = doc.add_heading(f"MATERIAL: {tema.upper()}", 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    lineas = limpiar_texto(texto_ia).split('\n')
    for linea in lineas:
        l = linea.strip()
        if not l: continue
        
        títulos_clave = ["#", "VERSIÓN", "GUION", "SCRIPT", "VOCABULARIO", "EJERCICIOS", "SOLUCIONARIO", "GLOSARIO"]
        if any(keyword in l.upper() for keyword in títulos_clave) and len(l) < 100:
            if "SOLUCIONARIO" in l.upper(): doc.add_page_break()
            level = 1 if l.startswith('#') else 2
            doc.add_heading(l.replace('#', '').strip(), level=level)
            continue

        p = doc.add_paragraph()
        partes = l.split('**')
        for i, parte in enumerate(partes):
            run = p.add_run(parte)
            if i % 2 == 1:
                run.bold = True
                run.font.color.rgb = RGBColor(200, 146, 74)
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- BARRA LATERAL ---
with st.sidebar:
    st.header("⚙️ Configuración")
    api_key = st.text_input("Gemini API Key", type="password")
    st.divider()
    logo_subido = st.file_uploader("Logo de tu escuela", type=["png", "jpg"])
    nombre_escuela = st.text_input("Proyecto", "PodcastELE")
    nombre_profe = st.text_input("Profesor", "Mario")
    idioma_apoyo = st.selectbox("Idioma de apoyo", ["Ninguno", "Inglés", "Polaco", "Francés"])

# --- INTERFAZ ---
st.title("🎙️ PodcastELE Pro: Dual (Blog + Podcast)")
col1, col2 = st.columns([2, 1])

with col1:
    tema_input = st.text_area("Tema/Idea", placeholder="Ej: Marga viaja a México...")
    instrucciones_extra = st.text_input("Extras", placeholder="Ej: Final abierto...")

with col2:
    nivel_mcer = st.selectbox("Nivel", ["A1", "A2", "B1", "B2", "C1", "C2"])
    duracion = st.select_slider("Duración", options=["5 min", "10 min", "15 min"])
    genero = st.selectbox("Género", ["Misterio", "Histórico", "Humor", "Fábula"])

# --- LÓGICA DE GENERACIÓN ---
if st.button("✨ Generar Material Completo"):
    if not api_key or not tema_input:
        st.warning("⚠️ Datos incompletos.")
    else:
        try:
            with st.spinner("Conectando con la API y redactando..."):
                # Intentamos listar modelos para encontrar el nombre exacto
                url_list = f"https://generativelanguage.googleapis.com/v1beta/models?key={api_key.strip()}"
                res_list = requests.get(url_list).json()
                
                # Buscamos modelos que soporten generación de contenido
                modelos = [m["name"] for m in res_list.get("models", []) if "generateContent" in m.get("supportedGenerationMethods", [])]
                
                # Lógica de selección inteligente de modelo
                modelo_final = ""
                opciones = ["models/gemini-1.5-pro", "models/gemini-1.5-flash", "models/gemini-pro"]
                
                for opcion in opciones:
                    if opcion in modelos:
                        modelo_final = opcion
                        break
                
                if not modelo_final and modelos:
                    modelo_final = modelos[0]
                
                if not modelo_final:
                    st.error("No se encontraron modelos compatibles en tu cuenta.")
                    st.stop()

                soporte = f"Usa el {idioma_apoyo} para traducciones." if idioma_apoyo != "Ninguno" else ""
                
                prompt = (
                    f"Eres editor ELE y guionista. Tema: {tema_input}. Nivel: {nivel_mcer}.\n"
                    f"ENTREGA ESTO:\n"
                    f"1. # VERSIÓN PARA EL BLOG (ALUMNO): Cuento narrativo literario.\n"
                    f"2. # VERSIÓN GUION (PODCAST): Guion con marcas [MÚSICA] y [SFX].\n"
                    f"3. # GLOSARIO Y EJERCICIOS: 10 términos y actividades.\n"
                    f"4. # SOLUCIONARIO\n"
                    f"{soporte} {instrucciones_extra}. Firma: {nombre_profe}."
                )

                # Construcción de la URL de generación (Aseguramos el nombre completo del modelo)
                url_gen = f"https://generativelanguage.googleapis.com/v1beta/{modelo_final}:generateContent?key={api_key.strip()}"
                
                res_gen = requests.post(url_gen, json={"contents": [{"parts": [{"text": prompt}]}]})
                
                if res_gen.status_code == 200:
                    st.session_state['material_podcast'] = res_gen.json()["candidates"][0]["content"]["parts"][0]["text"]
                    st.success(f"¡Contenido generado con {modelo_final}!")
                else:
                    st.error(f"Error API ({res_gen.status_code}): {res_gen.text}")
        except Exception as e:
            st.error(f"Error de conexión: {e}")

# --- VISUALIZACIÓN ---
if 'material_podcast' in st.session_state:
    st.divider()
    contenido = st.session_state['material_podcast']
    docx_bytes = generar_docx_podcast(contenido, nombre_escuela, nombre_profe, tema_input, nivel_mcer, logo_file=logo_subido)
    
    st.download_button("📥 Descargar Material (Word)", data=docx_bytes, file_name=f"PodcastELE_{nivel_mcer}.docx")

    t1, t2, t3 = st.tabs(["📖 Versión Alumno", "🎙️ Guion Podcast", "📝 Ejercicios"])
    
    with t1:
        if "# VERSIÓN PARA EL BLOG" in contenido:
            st.markdown(contenido.split("# VERSIÓN PARA EL BLOG")[1].split("# VERSIÓN GUION")[0])
        else: st.markdown(contenido)

    with t2:
        if "# VERSIÓN GUION" in contenido:
            st.markdown(contenido.split("# VERSIÓN GUION")[1].split("# GLOSARIO")[0])
        else: st.write("Sección de guion no detectada.")

    with t3:
        if "# GLOSARIO" in contenido:
            st.markdown("# GLOSARIO" + contenido.split("# GLOSARIO")[1])
